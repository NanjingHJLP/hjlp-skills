#!/usr/bin/env python3
"""
Word 文档导出脚本 - 将 Markdown 简历转换为 DOCX
用法: python export_docx.py <input.md> <output.docx> [--template <template_name>]
"""

import sys
import os
from pathlib import Path

try:
    import markdown
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误: 需要安装 markdown 和 python-docx")
    print("运行: pip install markdown python-docx")
    sys.exit(1)


def markdown_to_docx(markdown_content: str, output_path: str, template: str = "professional"):
    """将 Markdown 内容转换为 DOCX"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    lines = markdown_content.strip().split('\n')
    in_list = False

    for line in lines:
        stripped = line.strip()

        # 标题处理
        if stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=3)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 列表处理
        elif stripped.startswith('- ') or stripped.startswith('* '):
            if not in_list:
                in_list = True
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
        # 空行
        elif not stripped:
            in_list = False
            doc.add_paragraph()
        # 普通段落
        else:
            in_list = False
            # 处理加粗
            if '**' in stripped:
                parts = stripped.split('**')
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # 加粗部分
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(stripped)

    doc.save(output_path)
    print(f"DOCX 已生成: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python export_docx.py <input.md> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    template = "professional"

    input_path = Path(input_file)
    if not input_path.exists():
        print(f"错误: 输入文件不存在: {input_file}")
        sys.exit(1)

    content = input_path.read_text(encoding='utf-8')
    markdown_to_docx(content, output_file, template)
